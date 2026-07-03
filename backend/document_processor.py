"""
Document Processor for ThesisMind AI.
Handles PDF, DOCX, and TXT file processing.
"""

from pathlib import Path
from typing import Dict, List, Optional
import os


class DocumentProcessor:
    """Process academic documents."""
    
    SUPPORTED_FORMATS = ["pdf", "docx", "txt", "doc"]
    
    def __init__(self):
        """Initialize document processor."""
        pass
    
    def process_document(self, file_path: str) -> Optional[Dict]:
        """
        Process a document and extract information.
        
        Args:
            file_path: Path to the document
        
        Returns:
            Dict: Processed document information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return None
        
        file_ext = file_path.suffix.lower().strip(".")
        
        if file_ext == "pdf":
            return self._process_pdf(file_path)
        elif file_ext in ["docx", "doc"]:
            return self._process_docx(file_path)
        elif file_ext == "txt":
            return self._process_txt(file_path)
        
        return None
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """
        Process PDF file.
        
        Args:
            file_path: Path to PDF
        
        Returns:
            Dict: Document information
        """
        try:
            import PyPDF2
            
            text = ""
            pages = 0
            tables = []
            
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                pages = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return {
                "filename": file_path.name,
                "format": "pdf",
                "pages": pages,
                "text": text,
                "tables": tables,
                "metadata": {}
            }
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return None
    
    def _process_docx(self, file_path: Path) -> Dict:
        """
        Process DOCX file.
        
        Args:
            file_path: Path to DOCX
        
        Returns:
            Dict: Document information
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            tables = []
            pages = max(1, len(doc.paragraphs) // 50)  # Estimate pages
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            
            return {
                "filename": file_path.name,
                "format": "docx",
                "pages": pages,
                "text": text,
                "tables": tables,
                "metadata": {}
            }
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return None
    
    def _process_txt(self, file_path: Path) -> Dict:
        """
        Process TXT file.
        
        Args:
            file_path: Path to TXT
        
        Returns:
            Dict: Document information
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            pages = max(1, len(text) // 3000)  # Estimate pages
            
            return {
                "filename": file_path.name,
                "format": "txt",
                "pages": pages,
                "text": text,
                "tables": [],
                "metadata": {}
            }
        except Exception as e:
            print(f"Error processing TXT: {e}")
            return None
