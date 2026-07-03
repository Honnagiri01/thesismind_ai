"""
Export manager for ThesisMind AI.
Handles thesis export to multiple formats.
"""

from typing import Dict, BinaryIO
from pathlib import Path
from loguru import logger


class ExportManager:
    """Export generated thesis to various formats."""
    
    def __init__(self):
        """Initialize export manager."""
        self.supported_formats = ["docx", "pdf", "markdown", "latex"]
    
    def export_thesis(self, thesis: Dict, format_type: str, output_path: Path) -> bool:
        """
        Export thesis to specified format.
        
        Args:
            thesis: Generated thesis content
            format_type: Export format (docx, pdf, markdown, latex)
            output_path: Output file path
        
        Returns:
            bool: Success status
        """
        if format_type not in self.supported_formats:
            logger.error(f"Unsupported format: {format_type}")
            return False
        
        try:
            if format_type == "docx":
                return self._export_docx(thesis, output_path)
            elif format_type == "pdf":
                return self._export_pdf(thesis, output_path)
            elif format_type == "markdown":
                return self._export_markdown(thesis, output_path)
            elif format_type == "latex":
                return self._export_latex(thesis, output_path)
        except Exception as e:
            logger.error(f"Error exporting thesis: {e}")
            return False
    
    def _export_markdown(self, thesis: Dict, output_path: Path) -> bool:
        """
        Export thesis as Markdown.
        
        Args:
            thesis: Thesis content
            output_path: Output file path
        
        Returns:
            bool: Success status
        """
        try:
            content = f"# {thesis['title']}\n\n"
            content += f"**Type**: {thesis['type']}\n\n"
            content += "---\n\n"
            
            # Add table of contents
            content += "## Table of Contents\n\n"
            for i, chapter in enumerate(thesis["chapters"].keys(), 1):
                content += f"{i}. [{chapter}](#{chapter.lower().replace(' ', '-')})\n"
            
            content += "\n---\n\n"
            
            # Add chapters
            for chapter_title, chapter_content in thesis["chapters"].items():
                content += f"## {chapter_title}\n\n"
                content += chapter_content + "\n\n"
                content += "---\n\n"
            
            # Write to file
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Thesis exported to Markdown: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Error exporting to Markdown: {e}")
            return False
    
    def _export_docx(self, thesis: Dict, output_path: Path) -> bool:
        """
        Export thesis as Word document.
        
        Args:
            thesis: Thesis content
            output_path: Output file path
        
        Returns:
            bool: Success status
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(thesis["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add type
            type_para = doc.add_paragraph(f"Type: {thesis['type']}")
            type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Add table of contents
            doc.add_heading("Table of Contents", level=1)
            for i, chapter in enumerate(thesis["chapters"].keys(), 1):
                doc.add_paragraph(f"{i}. {chapter}", style="List Number")
            
            doc.add_page_break()
            
            # Add chapters
            for chapter_title, chapter_content in thesis["chapters"].items():
                doc.add_heading(chapter_title, level=1)
                doc.add_paragraph(chapter_content)
                doc.add_page_break()
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            logger.info(f"Thesis exported to DOCX: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return False
    
    def _export_pdf(self, thesis: Dict, output_path: Path) -> bool:
        """
        Export thesis as PDF.
        
        Args:
            thesis: Thesis content
            output_path: Output file path
        
        Returns:
            bool: Success status
        """
        try:
            md_path = output_path.with_suffix(".md")
            self._export_markdown(thesis, md_path)
            logger.warning("PDF export requires additional setup")
            return False
        
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            return False
    
    def _export_latex(self, thesis: Dict, output_path: Path) -> bool:
        """
        Export thesis as LaTeX.
        
        Args:
            thesis: Thesis content
            output_path: Output file path
        
        Returns:
            bool: Success status
        """
        try:
            content = "\\documentclass{report}\n\\usepackage[utf-8]{inputenc}\n\\usepackage{graphicx}\n\\usepackage{hyperref}\n\n"
            content += f"\\title{{{thesis['title']}}}\n\\author{{ThesisMind AI}}\n\\date{{\\today}}\n\n"
            content += "\\begin{document}\n\n\\maketitle\n\n\\tableofcontents\n\n"
            
            # Add chapters
            for chapter_title, chapter_content in thesis["chapters"].items():
                content += f"\\chapter{{{chapter_title}}}\n\n"
                escaped_content = chapter_content.replace("#", "\\#").replace("&", "\\&")
                content += escaped_content + "\n\n"
            
            content += "\\end{document}"
            
            # Write to file
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Thesis exported to LaTeX: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Error exporting to LaTeX: {e}")
            return False
